import os
from pathlib import Path
from typing import List, Tuple, Dict, Any
import docx
from docx import Document
from pypdf import PdfReader

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = ['.docx', '.one', '.pdf']
    
    def process_file(self, file_path: str, filename: str) -> Tuple[List[str], Dict[str, Any]]:
        """
        Process a file and return chunks and metadata.
        
        Args:
            file_path: Path to the file
            filename: Original filename
            
        Returns:
            Tuple of (chunks, metadata)        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.docx':
            return self._process_docx(file_path, filename)
        elif file_ext == '.one':
            return self._process_onenote(file_path, filename)
        elif file_ext == '.pdf':
            return self._process_pdf(file_path, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def _process_docx(self, file_path: str, filename: str) -> Tuple[List[str], Dict[str, Any]]:
        """Process Word document."""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # Combine all text
            full_text = "\n\n".join(text_content)
            
            # Simple chunking for now (can be improved later)
            chunks = self._chunk_text(full_text)
            
            metadata = {
                'filename': filename,
                'file_type': 'docx',
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'word_count': len(full_text.split())
            }
            
            return chunks, metadata
            
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def _process_onenote(self, file_path: str, filename: str) -> Tuple[List[str], Dict[str, Any]]:
        """
        Attempt to process a OneNote file. Local .one files are not supported for direct parsing.
        Suggest exporting to a supported format.
        """
        chunks = [
            f"OneNote file '{filename}' uploaded. Local .one files cannot be processed directly. ",
            "Please export your OneNote notebook to Word (.docx), PDF, or HTML and upload that file instead.\n",
            "See the documentation for more details."
        ]
        metadata = {
            'filename': filename,
            'file_type': 'onenote',
            'status': 'unsupported',
            'note': 'Local .one files are not supported. Please export to a supported format.'
        }
        return chunks, metadata
    
    def _process_pdf(self, file_path: str, filename: str) -> Tuple[List[str], Dict[str, Any]]:
        """Process PDF document."""
        try:
            reader = PdfReader(file_path)
            
            # Extract text from all pages
            text_content = []
            total_pages = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Add page number for reference
                        page_content = f"[Page {page_num + 1}]\n{page_text.strip()}"
                        text_content.append(page_content)
                except Exception as e:
                    # Skip problematic pages but continue processing
                    print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                    continue
            
            # Combine all text
            full_text = "\n\n".join(text_content)
            
            # Simple chunking
            chunks = self._chunk_text(full_text)
            
            metadata = {
                'filename': filename,
                'file_type': 'pdf',
                'total_pages': total_pages,
                'pages_processed': len(text_content),
                'word_count': len(full_text.split()) if full_text else 0
            }
            
            return chunks, metadata
            
        except Exception as e:
            raise Exception(f"Error processing PDF document: {str(e)}")
    
    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Simple text chunking with overlap.
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk in characters
            overlap: Overlap between chunks in characters
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings
                sentence_endings = ['.', '!', '?', '\n\n']
                best_break = end
                
                for i in range(end - 100, end + 100):
                    if i < len(text) and text[i] in sentence_endings:
                        best_break = i + 1
                        break
                
                end = best_break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            
            if start >= len(text):
                break
        
        return chunks
